"""
Purpose: Quote export generation — DOCX (python-docx), PDF (WeasyPrint), Supabase Storage upload.
         Export flow: build_snapshot → render → upload (outside DB txn) → return file_path.
Owner: [Claude]
"""
import io
from decimal import Decimal
from typing import Any

import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from supabase import create_client, Client

try:
    from weasyprint import HTML as _WeasyHTML
    _WEASYPRINT_AVAILABLE = True
except (OSError, ImportError):
    _WeasyHTML = None  # type: ignore[assignment]
    _WEASYPRINT_AVAILABLE = False

from app.config import settings
from app.constants import (
    EXPORT_PATH_ORIGINAL,
    EXPORT_PATH_REVISION,
    NEXTAN_WARRANTY_EXCLUSIONS,
)
from app.services.preview_service import render_html_preview


def _get_supabase() -> Client:
    """
    Purpose: Create a Supabase client using service-role key from settings.
    Inputs: none
    Outputs: supabase.Client
    Owner: [Claude]
    """
    return create_client(settings.supabase_url, settings.supabase_service_key)


def build_file_path(user_id: str, ref_number: str, file_type: str, revision_number: int) -> str:
    """
    Purpose: Construct the Supabase Storage path for an export file.
             Original (revision 0): exports/{user_id}/{ref}.{ext}
             Revisions: exports/{user_id}/{ref}-R{n}.{ext}
    Inputs: user_id (str), ref_number (str), file_type ('docx'|'pdf'), revision_number (int)
    Outputs: str — storage path
    Owner: [Claude]
    """
    ref_safe = ref_number.replace("/", "-")
    if revision_number == 0:
        return EXPORT_PATH_ORIGINAL.format(user_id=user_id, ref=ref_safe, ext=file_type)
    return EXPORT_PATH_REVISION.format(
        user_id=user_id, ref=ref_safe, revision=revision_number, ext=file_type
    )


def build_snapshot(
    sheet: Any,
    scenario: Any,
    line_items_with_pricing: list[dict],
    totals: Any,
    fx_rates_used: dict[str, Decimal],
    logo_url: str | None,
    signature_url: str | None,
    global_tnc: list[str],
    sheet_tnc: list[str],
) -> dict:
    """
    Purpose: Build the snapshot_data JSONB dict capturing full scenario state at export time.
             This is the source of truth for the locked quote. Stored in QuoteExport.snapshot_data.
    Inputs: sheet, scenario, line_items_with_pricing, totals, fx_rates_used,
            logo_url, signature_url, global_tnc, sheet_tnc
    Outputs: dict — JSON-serialisable snapshot
    Owner: [Claude]
    """
    def decimal_to_str(v):
        """Purpose: Convert Decimal to str for JSON serialisation. Owner: [Claude]"""
        return str(v) if isinstance(v, Decimal) else v

    return {
        "schema_version": 1,
        "sheet": {
            "id": str(sheet.id),
            "ref_number": sheet.ref_number,
            "date": str(sheet.date),
            "quote_title": sheet.quote_title,
            "client_name": sheet.client_name,
            "contact_name": sheet.contact_name,
            "contact_email": sheet.contact_email,
            "payment_term": sheet.payment_term,
            "quotation_validity_days": sheet.quotation_validity_days,
            "lead_time": sheet.lead_time,
            "local_tax": sheet.local_tax,
            "warranty": sheet.warranty,
            "general_notes": sheet.general_notes,
        },
        "scenario": {
            "id": str(scenario.id),
            "name": scenario.name,
            "discount_type": scenario.discount_type,
            "discount_value": decimal_to_str(scenario.discount_value),
            "show_gst": scenario.show_gst,
            "notes_exclusions": scenario.notes_exclusions or [],
        },
        "line_items": line_items_with_pricing,
        "totals": {
            "subtotal_sgd": decimal_to_str(totals.subtotal_sgd),
            "discount_amount_sgd": decimal_to_str(totals.discount_amount_sgd),
            "total_before_gst_sgd": decimal_to_str(totals.total_before_gst_sgd),
            "gst_amount_sgd": decimal_to_str(totals.gst_amount_sgd),
            "grand_total_sgd": decimal_to_str(totals.grand_total_sgd),
        },
        "fx_rates_used": {k: decimal_to_str(v) for k, v in fx_rates_used.items()},
        "logo_url": logo_url,
        "signature_url": signature_url,
        "global_tnc": global_tnc,
        "sheet_tnc": sheet_tnc,
        "warranty_exclusions": NEXTAN_WARRANTY_EXCLUSIONS,
    }


def render_pdf(context: dict) -> bytes:
    """
    Purpose: Render the quote as a PDF using WeasyPrint.
             Uses the shared Jinja2 HTML template from preview_service.
    Inputs: context (dict) — same shape as render_html_preview
    Outputs: bytes — PDF binary content
    Owner: [Claude]
    """
    if not _WEASYPRINT_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="PDF export requires GTK libraries not available on this platform. Use DOCX export or deploy to Linux.",
        )
    html_content = render_html_preview(context)
    return _WeasyHTML(string=html_content).write_pdf()


# ---------------------------------------------------------------------------
# DOCX rendering helpers
# ---------------------------------------------------------------------------

def _fetch_image(url: str) -> io.BytesIO | None:
    """
    Purpose: Fetch an image from a URL and return as BytesIO for embedding in DOCX.
    Inputs: url (str)
    Outputs: io.BytesIO | None — None on any error
    Owner: [Claude]
    """
    try:
        resp = httpx.get(url, timeout=10.0)
        resp.raise_for_status()
        buf = io.BytesIO(resp.content)
        buf.seek(0)
        return buf
    except Exception:
        return None


def _set_cell_bg(cell, hex_color: str) -> None:
    """
    Purpose: Set a table cell background colour via direct XML manipulation.
             python-docx does not expose this through its public API.
    Inputs: cell (docx.table._Cell), hex_color (str, e.g. 'D9D9D9')
    Outputs: none — mutates cell XML in place
    Owner: [Claude]
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_cell(cell, text: str, align: str = "left") -> None:
    """
    Purpose: Write bold text into a table cell with optional alignment.
    Inputs: cell, text (str), align ('left'|'right'|'center')
    Outputs: none — mutates cell in place
    Owner: [Claude]
    """
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = True
    if align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _right_cell(cell, text: str) -> None:
    """
    Purpose: Write right-aligned text into a table cell.
    Inputs: cell, text (str)
    Outputs: none — mutates cell in place
    Owner: [Claude]
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _remove_table_borders(table) -> None:
    """
    Purpose: Remove all borders from a table (invisible table for layout only).
    Inputs: table (docx.table.Table)
    Outputs: none — mutates table XML in place
    Owner: [Claude]
    """
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _fmt(val) -> str:
    """
    Purpose: Format a numeric value as a comma-separated 2dp string, or '—' if zero/None.
    Inputs: val (str | Decimal | None)
    Outputs: str
    Owner: [Claude]
    """
    try:
        n = Decimal(str(val))
        return f"{n:,.2f}"
    except Exception:
        return "—"


# ---------------------------------------------------------------------------
# Main DOCX renderer
# ---------------------------------------------------------------------------

SECTIONS_ORDER = ["Hardware", "Software", "Professional Fees", "Maintenance"]

_FONT = "Times New Roman"
_FONT_SIZE = Pt(12)


def _set_font(run, bold: bool = False, size=None) -> None:
    """
    Purpose: Apply Times New Roman font to a run, matching original NEXTAN sample.
    Inputs: run (docx Run), bold (bool), size (Pt | None — defaults to 12pt)
    Outputs: none — mutates run in place
    Owner: [Claude]
    """
    run.font.name = _FONT
    run.font.size = size if size is not None else _FONT_SIZE
    run.font.bold = bold


def _cell_write(cell, text: str, bold: bool = False, size=None,
                align: WD_ALIGN_PARAGRAPH | None = None):
    """
    Purpose: Clear a cell and write a single run with Times New Roman formatting.
    Inputs: cell, text (str), bold (bool), size (Pt|None), align (WD_ALIGN_PARAGRAPH|None)
    Outputs: the paragraph — mutates cell in place
    Owner: [Claude]
    """
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    _set_font(run, bold=bold, size=size)
    return para


def render_docx(context: dict) -> bytes:
    """
    Purpose: Render the quote as a branded NEXTAN DOCX quotation.
             Matches the original sample exactly:
             Word page header (logo + address) → client header table → quote title →
             main items table (no section rows) → asterisk notes → maintenance table →
             Terms & Conditions heading → terms table → warranty T&C → closing → sign-off.
    Inputs: context (dict) — output of _build_export_context in exports.py
    Outputs: bytes — DOCX binary content
    Owner: [Claude]
    """
    sheet = context["sheet"]
    line_items = context.get("line_items", [])
    totals = context["totals"]
    show_gst = context.get("show_gst", False)
    warranty_exclusions = context.get("warranty_exclusions", NEXTAN_WARRANTY_EXCLUSIONS)
    global_tnc = context.get("global_tnc", [])
    sheet_tnc = context.get("sheet_tnc", [])
    notes_exclusions = context.get("notes_exclusions", [])
    logo_url = context.get("logo_url")
    signature_url = context.get("signature_url")
    company_name = context.get("company_name", "NEXTAN Pte Ltd")
    company_contact_name = context.get("company_contact_name", "Justin Low")
    company_contact_email = context.get("company_contact_email", "")
    company_contact_phone = context.get("company_contact_phone", "")

    doc = Document()

    # A4 — 1.0" margins all sides, 0.125" header distance (matches original)
    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.header_distance = Inches(0.125)

    # ------------------------------------------------------------------ #
    # PAGE HEADER — 2-col borderless table: logo (left) | address (right)
    # Matches original: r0c0=2.575" logo image, r0c1=3.725" company text
    # ------------------------------------------------------------------ #
    page_hdr = doc.sections[0].header
    # Remove default empty paragraph that Word adds to headers
    for p in page_hdr.paragraphs:
        p.clear()

    hdr_tbl = page_hdr.add_table(rows=1, cols=2, width=Inches(6.27))
    _remove_table_borders(hdr_tbl)
    hdr_tbl.cell(0, 0).width = Inches(2.575)
    hdr_tbl.cell(0, 1).width = Inches(3.725)

    # Left col — logo
    if logo_url:
        logo_bytes = _fetch_image(logo_url)
        if logo_bytes:
            logo_cell = hdr_tbl.cell(0, 0)
            logo_cell.text = ""
            logo_para = logo_cell.paragraphs[0]
            logo_para.add_run().add_picture(logo_bytes, width=Inches(2.42))

    # Right col — company address block
    addr_cell = hdr_tbl.cell(0, 1)
    addr_cell.text = ""
    addr_lines = [company_name]
    for i, line in enumerate(addr_lines):
        if i == 0:
            para = addr_cell.paragraphs[0]
        else:
            para = addr_cell.add_paragraph()
        _set_font(para.add_run(line), bold=(i == 0), size=Pt(10))

    # ------------------------------------------------------------------ #
    # 1. CLIENT HEADER TABLE — 4r x 5c, no borders
    # Col widths matching original: 0.56", 0.19", 3.58", 0.75", 1.78"
    # ------------------------------------------------------------------ #
    client_tbl = doc.add_table(rows=4, cols=5)
    _remove_table_borders(client_tbl)
    col_w = [Inches(0.56), Inches(0.19), Inches(3.58), Inches(0.75), Inches(1.78)]
    for row in client_tbl.rows:
        for ci, w in enumerate(col_w):
            row.cells[ci].width = w

    # Row 0: Date | : | <date> | Ref No.: | <ref>
    _cell_write(client_tbl.cell(0, 0), "Date")
    _cell_write(client_tbl.cell(0, 1), ":")
    _cell_write(client_tbl.cell(0, 2), sheet.get("date") or "")
    _cell_write(client_tbl.cell(0, 3), "Ref No.:")
    _cell_write(client_tbl.cell(0, 4), sheet.get("ref_number") or "")

    # Rows 1-3: To / Attn. / Email — cols 2-4 merged
    for ri, (label, key) in enumerate([
        ("To", "client_name"),
        ("Attn.", "contact_name"),
        ("Email", "contact_email"),
    ], start=1):
        _cell_write(client_tbl.cell(ri, 0), label)
        _cell_write(client_tbl.cell(ri, 1), ":")
        merged = client_tbl.cell(ri, 2).merge(client_tbl.cell(ri, 4))
        _cell_write(merged, sheet.get(key) or "")

    # ------------------------------------------------------------------ #
    # 2. QUOTE TITLE — Times New Roman 12pt bold
    # ------------------------------------------------------------------ #
    title_para = doc.add_paragraph()
    _set_font(title_para.add_run(f"Quotation: {sheet.get('quote_title') or ''}"),
              bold=True, size=Pt(12))

    # ------------------------------------------------------------------ #
    # 3. LINE ITEMS TABLE — Table Grid, no section header rows
    # Col widths matching original: 0.39", 3.37", 0.70", 1.15", 1.18"
    # No. column header present but item cells left blank (matches original)
    # ------------------------------------------------------------------ #
    item_col_w = [Inches(0.39), Inches(3.37), Inches(0.70), Inches(1.15), Inches(1.18)]

    # Separate visible top-level items: main sections vs Maintenance
    main_items: list[dict] = []
    maint_items: list[dict] = []
    for item in line_items:
        if not item.get("is_visible", True):
            continue
        if item.get("parent_line_item_id"):
            continue
        sec = item.get("section", "Hardware")
        (maint_items if sec == "Maintenance" else main_items).append(item)

    # Sort within each group by section order then display_order
    sec_rank = {s: i for i, s in enumerate(SECTIONS_ORDER)}
    main_items.sort(key=lambda x: (
        sec_rank.get(x.get("section", "Hardware"), 99),
        x.get("display_order", 0),
    ))
    maint_items.sort(key=lambda x: x.get("display_order", 0))

    def _build_items_table(items: list[dict]) -> "docx.table.Table":
        """
        Purpose: Build a 5-col Table Grid matching the original NEXTAN format.
                 Header row then one row per item. No section separator rows.
                 No. column header present; item cells left blank per original.
        Inputs: items (list[dict]) — visible top-level line items
        Outputs: docx Table — added to doc, returned for caller to append total rows
        Owner: [Claude]
        """
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        for ci, w in enumerate(item_col_w):
            tbl.rows[0].cells[ci].width = w

        # Header row — bold, right-aligned for numeric cols
        for ci, hdr_text in enumerate(
            ["No", "Description", "Qty", "Unit Price\n(S$)", "Total Price\n(S$)"]
        ):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            if ci >= 2:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_font(para.add_run(hdr_text), bold=True)

        for item in items:
            computed = item.get("computed", {})
            line_total = Decimal(str(computed.get("line_total_sgd") or "0"))
            selling_rate = computed.get("selling_rate_sgd")

            row = tbl.add_row()
            cells = row.cells

            # No. — blank to match original sample
            _cell_write(cells[0], "")

            # Description + sub_specs on separate lines
            cells[1].text = ""
            desc_para = cells[1].paragraphs[0]
            _set_font(desc_para.add_run(item.get("description", "")))
            sub_specs = item.get("sub_specs") or []
            if isinstance(sub_specs, list):
                for spec in sub_specs:
                    sp = cells[1].add_paragraph()
                    _set_font(sp.add_run(str(spec)))

            # Qty — right-aligned
            cells[2].text = ""
            qty_para = cells[2].paragraphs[0]
            qty_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_font(qty_para.add_run(str(item.get("qty", 1))))

            # Unit Price — right-aligned
            cells[3].text = ""
            up_para = cells[3].paragraphs[0]
            up_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_font(up_para.add_run(_fmt(selling_rate) if selling_rate else "—"))

            # Total Price — right-aligned
            cells[4].text = ""
            tp_para = cells[4].paragraphs[0]
            tp_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_font(tp_para.add_run(_fmt(line_total) if line_total else "—"))

        return tbl

    # Build main items table and append totals rows
    if main_items:
        main_tbl = _build_items_table(main_items)

        subtotal = Decimal(str(totals.get("subtotal_sgd", "0")))
        discount = Decimal(str(totals.get("discount_amount_sgd", "0")))
        gst_amount = Decimal(str(totals.get("gst_amount_sgd", "0")))
        grand_total = Decimal(str(totals.get("grand_total_sgd", "0")))

        def _total_row(tbl, label: str, value: str, bold: bool = False) -> None:
            """Purpose: Append a totals row — label merged cols 0-3, value in col 4. Owner: [Claude]"""
            row = tbl.add_row()
            merged = row.cells[0].merge(row.cells[3])
            _cell_write(merged, label, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _cell_write(row.cells[4], value, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)

        if discount > 0:
            _total_row(main_tbl, "Subtotal (SGD)", f"$ {_fmt(subtotal)}")
            _total_row(main_tbl, "Discount (SGD)", f"- {_fmt(discount)}")

        if show_gst:
            _total_row(main_tbl, "GST (9%)", f"$ {_fmt(gst_amount)}")

        # TOTAL row — bold, "TOTAL" label + "$ X,XXX.XX" matching original
        _total_row(main_tbl, "TOTAL", f"$ {_fmt(grand_total)}", bold=True)

    # ------------------------------------------------------------------ #
    # 4. ASTERISK NOTES — paragraphs after main table, before maintenance
    # ------------------------------------------------------------------ #
    for note in (notes_exclusions or []):
        note_para = doc.add_paragraph()
        _set_font(note_para.add_run(str(note)))

    # ------------------------------------------------------------------ #
    # 5. MAINTENANCE TABLE — separate table with own header if items exist
    # ------------------------------------------------------------------ #
    if maint_items:
        maint_heading = doc.add_paragraph()
        _set_font(maint_heading.add_run("Annual Maintenance"))
        _build_items_table(maint_items)

    # ------------------------------------------------------------------ #
    # 6. TERMS & CONDITIONS HEADING (before the terms table — matches original)
    # ------------------------------------------------------------------ #
    tnc_heading = doc.add_paragraph()
    _set_font(tnc_heading.add_run("Terms & Conditions:"), bold=True)

    # ------------------------------------------------------------------ #
    # 7. TERMS TABLE — 3 cols, no borders
    # Col widths: 1.45", 0.19", 4.74" (matching original)
    # ------------------------------------------------------------------ #
    validity_text = (
        f"{sheet.get('quotation_validity_days', 90)} days from the date of quotation."
    )
    terms_rows_data = [
        ("Payment Term", sheet.get("payment_term") or "TBD"),
        ("Quotation Validity", validity_text),
        ("Lead Time", sheet.get("lead_time") or "30 working days from signoff"),
        ("Local Tax", sheet.get("local_tax") or "Prices quoted in SGD are subjected to Singapore GST"),
        ("Warranty", sheet.get("warranty") or "Standard 12 month against manufacturing defects"),
    ]
    if notes_exclusions:
        terms_rows_data.append(("Note(s)", "\n".join(str(n) for n in notes_exclusions)))

    terms_tbl = doc.add_table(rows=len(terms_rows_data), cols=3)
    _remove_table_borders(terms_tbl)
    for row in terms_tbl.rows:
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(0.19)
        row.cells[2].width = Inches(4.74)

    for ri, (label, value) in enumerate(terms_rows_data):
        _cell_write(terms_tbl.rows[ri].cells[0], label, bold=True)
        _cell_write(terms_tbl.rows[ri].cells[1], ":")
        _cell_write(terms_tbl.rows[ri].cells[2], value)

    # ------------------------------------------------------------------ #
    # 8. WARRANTY T&C SECTION
    # ------------------------------------------------------------------ #
    warranty_para = doc.add_paragraph()
    _set_font(warranty_para.add_run("Warranty Terms:"), bold=True)

    excl_para = doc.add_paragraph()
    _set_font(excl_para.add_run("Exclusions from Warranty:"), bold=True)

    liability_para = doc.add_paragraph()
    liability_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(liability_para.add_run(
        "NEXTAN assumes no liability as a consequence of following circumstances, "
        "under which will be automatically excluded for warranty:"
    ))

    for excl in (warranty_exclusions or []):
        p = doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p.add_run(excl))

    for tnc_item in ((global_tnc or []) + (sheet_tnc or [])):
        p = doc.add_paragraph(style="List Paragraph")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_font(p.add_run(tnc_item))

    # ------------------------------------------------------------------ #
    # 9. CLOSING PARAGRAPH
    # ------------------------------------------------------------------ #
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_font(closing.add_run(
        "We trust that the above price is acceptable to you and look forward to receiving "
        "your favourable reply soon. Should you have any further queries, please do not "
        "hesitate to contact us."
    ))

    # ------------------------------------------------------------------ #
    # 10. SIGN-OFF TABLE — 4r x 3c, no borders
    # Col widths: 2.422", 1.053", 1.053" (matching original)
    # Signature image embedded in r0c0 below "Yours faithfully" (matches original)
    # ------------------------------------------------------------------ #
    signoff_tbl = doc.add_table(rows=4, cols=3)
    _remove_table_borders(signoff_tbl)
    for row in signoff_tbl.rows:
        row.cells[0].width = Inches(2.422)
        row.cells[1].width = Inches(1.053)
        row.cells[2].width = Inches(1.053)

    # r0c0: "Yours faithfully" + signature image below
    r0c0 = signoff_tbl.rows[0].cells[0]
    r0c0.text = ""
    yf_para = r0c0.paragraphs[0]
    _set_font(yf_para.add_run("Yours faithfully"))
    if signature_url:
        sig_bytes = _fetch_image(signature_url)
        if sig_bytes:
            sig_para = r0c0.add_paragraph()
            sig_para.add_run().add_picture(sig_bytes, width=Inches(0.94))

    # r1: name, r2: email, r3: phone
    _cell_write(signoff_tbl.rows[1].cells[0], company_contact_name or "")
    _cell_write(signoff_tbl.rows[2].cells[0],
                f"Email\t: {company_contact_email}" if company_contact_email else "")
    _cell_write(signoff_tbl.rows[3].cells[0],
                f"Phone\t: {company_contact_phone}" if company_contact_phone else "")

    # ------------------------------------------------------------------ #
    # Save to bytes
    # ------------------------------------------------------------------ #
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Storage helpers
# ---------------------------------------------------------------------------

async def upload_to_storage(file_bytes: bytes, file_path: str, file_type: str) -> str:
    """
    Purpose: Upload file bytes to Supabase Storage exports bucket.
             Called OUTSIDE the DB transaction to avoid holding DB connections during upload.
    Inputs: file_bytes (bytes), file_path (str), file_type ('docx'|'pdf')
    Outputs: str — the file_path (same as input, confirmed by Supabase)
    Owner: [Claude]
    """
    content_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if file_type == "docx"
        else "application/pdf"
    )
    supabase = _get_supabase()
    try:
        supabase.storage.from_(settings.supabase_storage_bucket_exports).upload(
            path=file_path,
            file=file_bytes,
            file_options={"content-type": content_type, "upsert": "true"},
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Export file upload failed: {exc}. No database record was created.",
        ) from exc
    return file_path


def generate_signed_url(file_path: str, expires_in: int = 3600) -> str:
    """
    Purpose: Generate a time-limited signed URL for a Supabase Storage file.
    Inputs: file_path (str), expires_in (int, seconds, default 3600)
    Outputs: str — signed URL
    Owner: [Claude]
    """
    supabase = _get_supabase()
    result = supabase.storage.from_(settings.supabase_storage_bucket_exports).create_signed_url(
        path=file_path, expires_in=expires_in
    )
    return result["signedURL"]


def delete_from_storage(file_path: str) -> None:
    """
    Purpose: Delete a file from Supabase Storage exports bucket.
             Called when a QuoteExport record is deleted.
    Inputs: file_path (str)
    Outputs: none
    Owner: [Claude]
    """
    supabase = _get_supabase()
    try:
        supabase.storage.from_(settings.supabase_storage_bucket_exports).remove([file_path])
    except Exception:
        # Non-fatal: log but don't raise. The DB record deletion is the authoritative action.
        pass
